from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'D:/AManager/pyProject/PhantomTower/样片工厂-首页功能确认说明.docx'
doc = Document(); sec = doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(.78)

for n,s,c,b,a in [('Normal',10.5,'28332F',0,6),('Heading 1',16,'266E88',15,7),('Heading 2',12.5,'1F5870',10,5)]:
    st=doc.styles[n]; st.font.name='Calibri'; st._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(s); st.font.color.rgb=RGBColor.from_string(c); st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a)
    if n != 'Normal': st.font.bold=True
doc.styles['Normal'].paragraph_format.line_spacing=1.15

def fill(cell,color):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); cell._tc.get_or_add_tcPr().append(shd)
def put(cell,text,bold=False,color=None):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(text); r.bold=bold; r.font.name='Calibri'; r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(9.2)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def tbl(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): t.columns[i].width=Inches(widths[i]); put(t.rows[0].cells[i],h,True,'FFFFFF'); fill(t.rows[0].cells[i],'266E88')
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].width=Inches(widths[i]); put(cs[i],v)
def bullet(text,check=False):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(('□  ' if check else '')+text)
def sub(title,text):
    doc.add_heading(title,2); doc.add_paragraph(text)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('样片工厂\n首页功能确认说明'); r.bold=True; r.font.size=Pt(23); r.font.color.rgb=RGBColor(24,65,83); r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(14); r=p.add_run('用于确认首页能力、操作流程、生成结果及提示词管理方式'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(92,105,101)
p=doc.add_paragraph(); x=p.add_run('本次更新：'); x.bold=True; x.font.color.rgb=RGBColor(38,110,136); p.add_run('产品名称统一调整为“样片工厂”；首页同时保留文字生图和图生图；内置提示词可以在页面调整并保存，无需为修改提示词重新部署程序。')

doc.add_heading('一、产品定位',1)
doc.add_paragraph('样片工厂是一个用于快速生成、筛选和继续编辑图片样片的工作台。用户先选择生成模式，再填写提示词或上传素材；系统按明确的素材角色生成多张候选样片。用户可以从结果中选择一张，继续进行局部调整。')

doc.add_heading('二、首页页面布局与操作顺序',1)
tbl(['页面区域','内容与操作','显示规则'],[
('顶部标题区','显示“样片工厂”、当前模式、开始生成与任务进度。','始终显示；生成时锁定本次任务所需配置。'),
('左侧模式与素材区','通过模式按钮切换文字生图或图生图；图生图时展示人物、目标图、背景、道具等上传项。','不同模式只展示当前需要的输入项，避免表单堆叠。'),
('中间结果区','显示预计样片数量、生成队列、结果分组、查看、重试、选择、导出和继续编辑。','始终作为工作重点；结果按任务或版本归类。'),
('右侧生成设置区','提示词、模型、版本、分辨率、尺寸、每项生成数量，以及提示词模板管理入口。','文字生图与图生图共用模型和规格设置；素材相关设置仅图生图显示。')],[1.35,2.85,2.3])
doc.add_paragraph('推荐交互顺序：选择模式 → 填写提示词或上传素材 → 选择模型、版本、分辨率、尺寸和数量 → 查看预计结果数量 → 生成样片 → 选择结果继续编辑。')

doc.add_heading('三、模式选择：文字生图与图生图',1)
tbl(['模式','用户输入','系统处理','生成结果'],[
('文字生图','只填写图像描述、选择模型与规格。','直接将提示词发送给图像生成接口，不要求上传图片。','按“每项生成数量”得到多张同主题候选样片。'),
('图生图','上传人物、主目标图、背景、道具等素材，可补充文字要求。','将上传素材按角色发送，并结合内置提示词执行人物替换、背景替换、道具替换或多图融合。','得到保留构图或按素材融合后的候选样片。')],[1.1,2.0,2.3,1.1])
sub('1. 文字生图页面','左侧仅保留模式选择；右侧突出“图像描述”输入框、可选提示词模板、模型版本、分辨率、尺寸和数量。参考图中的“智能文生图 / 智能图生图”卡片式切换可用于此入口。')
sub('2. 图生图页面','切换到图生图后，左侧展开素材工作台。用户上传图片并指定每张图片的角色；右侧继续保留提示词、模型和规格配置。提示词只用于补充意图，核心画面关系由素材角色和内置模板约束。')

doc.add_heading('四、图生图素材与处理方式',1)
tbl(['素材/功能','用户可以做什么','系统产生的结果'],[
('人物参考','上传 1-3 张人物图片。','固定人物身份、脸部、发型、服装和外观，是图生图的主体来源。'),
('逐张批处理目标图','上传多张目标图并选择逐张处理。','每张目标图独立形成任务，保留原构图、姿势、景别、背景关系和光线。'),
('主目标图 / 多图融合','选择一张主目标图，再添加人物、背景、道具。','主目标图决定画布结构，其余素材融合到同一画面，生成一组候选样片。'),
('背景参考','上传一张背景照片。','只替换人物背后的环境，保留人物、前景遮挡、比例、位置和镜头关系。'),
('道具参考','上传花瓶等道具，并指定被替换对象。','只替换画面指定道具，并匹配原位置、比例、透视、接触关系和光线。'),
('继续编辑','选中已有样片，框选或指定局部区域。','只修改背景、道具、服装、表情、手脚或光线等局部，生成新版本。')],[1.38,2.52,2.6])

doc.add_heading('五、内置提示词的可编辑与保存机制',1)
doc.add_paragraph('现有内置提示词不再写死在前端代码中。首页提供“提示词模板”管理入口，管理员可以查看、修改、保存和恢复模板。保存后新任务直接读取最新模板，不需要为提示词调整重新修改代码或重新部署程序。')
tbl(['模板类型','用途','管理员可调整内容'],[
('文字生图基础模板','为文字描述增加通用画质、构图或风格规则。','默认正向提示词、默认负向提示词、默认补充词。'),
('人物替换模板','约束人物替换时保留目标图姿势、构图和环境。','人物身份规则、保留项、禁止项、质量检查词。'),
('背景替换模板','约束仅替换背景，不重新生成主体。','可替换范围、必须保留范围、背景融合规则。'),
('道具替换模板','约束只替换用户指定的对象。','对象定位语句、透视比例规则、禁止生成额外道具。'),
('多图融合模板','约束主目标图、人物、背景和道具的角色边界。','各素材角色说明、融合优先级、冲突处理规则。'),
('局部编辑模板','约束只编辑蒙版或指定区域。','局部保留规则、局部修改规则、版本生成规则。')],[1.38,2.12,3.0])
sub('保存位置与 JSON 建议','提示词配置保存到 data 目录，例如 data/prompt-templates.json。文件中使用稳定的模板 ID、名称、适用模式、系统提示词、默认负向提示词、更新时间和版本号。程序启动时读取该 JSON；管理员在页面保存后通过后端接口更新 JSON，并返回最新版本。')
sub('管理规则','普通创作用户只选择模板和填写补充提示词；具备管理权限的用户才能编辑、保存、复制、恢复默认模板或查看历史版本。保存前显示修改预览，避免误改影响全部生成任务。')

doc.add_heading('六、规格设置与样片数量',1)
doc.add_paragraph('右侧规格区参考现有界面结构：模型模式切换、模型版本、分辨率、尺寸和每项生成数量。尺寸支持自适应、常用比例和自定义宽高；图生图的“自适应”根据主目标图或当前编辑基础图确定。')
tbl(['使用场景','数量计算','示例'],[
('文字生图','1 个文字任务 × 每项生成数量','提示词 1 条，数量 4 = 4 张样片'),
('图生图：逐张批处理','目标图数量 × 每张生成数量','3 张目标图 × 2 张 = 6 张样片'),
('图生图：多图融合','1 个融合任务 × 每项生成数量','人物 + 主目标图 + 背景 + 花瓶，数量 4 = 4 张样片'),
('局部继续编辑','选中样片数量 × 编辑生成数量','选中 1 张，生成 3 个局部版本 = 3 张')],[1.62,2.45,2.43])

doc.add_heading('七、最终可交付的使用结果',1)
for x in ['通过一句或多句文字描述直接生成图片样片。','通过人物参考图批量替换多张目标图中的人物。','在一张主目标图中融合人物、背景和指定道具。','用上传的背景照片替换背景，用上传的花瓶等图片替换指定道具。','根据设置一次生成多张候选样片，并在生成前看到预计数量。','从结果中选择满意版本，针对局部进行继续编辑，减少复杂提示词造成的全图错误。','管理员通过页面维护内置提示词模板，保存至 data 目录 JSON 后立即用于后续任务。','对样片进行查看、重试、筛选、导出和版本延续。']: bullet(x)

doc.add_heading('八、客户确认项',1)
doc.add_paragraph('请确认以下理解是否符合实际使用需求：')
for x in ['产品首页名称调整为“样片工厂”。','首页保留文字生图和图生图两种模式，并通过明显按钮或卡片切换。','文字生图不要求上传图片，直接使用图像描述调用接口。','图生图保留人物替换、逐张批处理、多图融合、背景替换、道具替换和局部继续编辑。','背景和道具参考图需要作为同一融合任务中的素材角色，而不是默认分别拆成不同任务。','道具替换时由用户指定需要替换的画面对象。','内置提示词可在管理页面调整，并保存到 data 目录下的 JSON 文件。','提示词修改不需要修改前端代码或重新部署；新任务读取保存后的最新模板。','样片数量根据文字任务、批处理任务、融合任务或局部编辑任务分别计算。']: bullet(x,True)

f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=f.add_run('样片工厂 · 首页功能确认说明'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(120,132,128)
doc.save(OUT); print(OUT)
