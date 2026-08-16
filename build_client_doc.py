from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

out = 'D:/AManager/pyProject/PhantomTower/首页图像创作工作台-功能确认说明.docx'
doc = Document(); sec = doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(.78)
styles=doc.styles
for name,size,color,before,after in [('Normal',10.5,'262D2B',0,6),('Heading 1',16,'2E74B5',14,7),('Heading 2',13,'2E74B5',11,5)]:
    s=styles[name]; s.font.name='Calibri'; s._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    if name!='Normal': s.font.bold=True
styles['Normal'].paragraph_format.line_spacing=1.12

def shade(c, fill):
    e=OxmlElement('w:shd'); e.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(e)
def cell(c, text, bold=False, color=None):
    c.text=''; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(text); r.bold=bold; r.font.name='Calibri'; r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(9.2)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def grid(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): t.columns[i].width=Inches(widths[i]); cell(t.rows[0].cells[i],h,True,'FFFFFF'); shade(t.rows[0].cells[i],'2E74B5')
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row): cs[i].width=Inches(widths[i]); cell(cs[i],val)
    return t
def b(text, check=False):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(('□  ' if check else '')+text)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('首页图像创作工作台\n功能确认说明'); r.bold=True; r.font.size=Pt(23); r.font.color.rgb=RGBColor(25,52,70); r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),'Microsoft YaHei')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(14); r=p.add_run('用于确认首页功能范围、操作方式与预期生成结果'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(95,105,103)
p=doc.add_paragraph(); r=p.add_run('确认重点：'); r.bold=True; r.font.color.rgb=RGBColor(46,116,181); p.add_run('首页是“人物参考 + 主目标图 + 背景/道具融合 + 批量样片 + 局部继续编辑”的图像创作工作台。')

doc.add_heading('一、首页整体定位',1)
doc.add_paragraph('用户以一个固定人物为核心，上传一张或多张参考素材，选择处理模式后生成可比较、可筛选、可继续编辑的样片。系统需要清楚区分“分别处理多张图片”和“多张素材共同融合到一张图片”两种工作方式。')
doc.add_heading('二、首页主要功能',1)
grid(['功能模块','用户可以做什么','系统会产生什么'],[
('人物参考','上传 1-3 张人物图片，确定主体人物。','锁定人物身份、脸部、发型、服装和外观；所有结果以此人物为基础。'),
('逐张批处理','上传多张目标图，要求系统逐张处理。','每张目标图形成独立任务，保留各自构图、姿势、背景和光线。'),
('多图融合','选择一张主目标图，并叠加人物、背景、道具等参考图。','所有素材共同生成同一画面的多张候选结果。'),
('背景替换','上传背景照片。','保留人物主体和构图，仅替换人物背后的环境。'),
('道具替换','上传花瓶等道具，并指定画面中的替换对象。','只替换指定道具，保留人物、背景、其他物件和光线关系。'),
('局部继续编辑','从生成结果中选一张，对指定区域进行调整。','以当前结果为基础进行局部修改，避免重新处理全部素材。'),
('数量控制','选择每个任务需要生成的样片数量。','按任务数量计算并生成对应数量的结果图。')],[1.25,2.35,2.9])
doc.add_heading('三、两种目标图处理方式',1)
doc.add_heading('1. 逐张批处理',2); doc.add_paragraph('适用于需要一次处理多张目标图的场景。每张目标图彼此独立，系统不会把多张目标图合成一张画面。')
for x in ['上传 1 张人物参考图 + 3 张目标图，每张生成 2 张，最终得到 6 张结果图。','每张结果保留对应目标图的构图、景别、人物位置、姿势、背景关系和光线。','结果区按目标图分组，方便查看与重试。']: b(x)
doc.add_heading('2. 多图融合生成',2); doc.add_paragraph('适用于将人物、背景、道具等多张参考素材共同组合到同一张画面中。必须指定一张主目标图作为最终画布。')
for x in ['主目标图决定画布、构图、景别、人物位置、镜头角度和光线。','人物参考图只决定人物身份与外观。','背景参考图只决定背景环境。','道具参考图只决定指定道具。','所有素材共同形成一个融合任务，再根据数量生成多张候选样片。']: b(x)
doc.add_heading('四、素材规则与内置提示词',1)
for h,t in [('人物替换','使用人物参考图中的人物替换目标图原人物，同时保留目标图的构图、姿势、景别、背景、遮挡和光线，不保留原人物残留肢体、衣物或影子。'),('背景替换','仅替换人物背后的环境，保留人物身份、姿势、位置、比例、前景遮挡、光线方向和镜头关系。'),('道具替换','仅替换用户指定的道具，并匹配原位置、比例、透视、接触关系和光线，不生成额外道具。用户需确认替换对象，例如“画面右侧桌上的花瓶”。'),('多图融合','主目标图决定画面结构，人物、背景、道具分别承担对应角色，所有素材融合到同一画面，不互相替换角色或背景。')]:
    doc.add_heading(h,2); doc.add_paragraph('系统自动要求：'+t)
doc.add_heading('五、生成结果与数量计算',1)
grid(['场景','数量计算方式','示例结果'],[('逐张批处理','目标图数量 × 每张生成数量','3 张目标图 × 2 张 = 6 张'),('多图融合','1 个融合任务 × 每项生成数量','人物 + 目标图 + 背景 + 道具，生成 4 张 = 4 张'),('局部编辑','选中样片数量 × 编辑生成数量','选中 1 张，生成 3 个局部版本 = 3 张')],[1.35,2.5,2.65])
doc.add_paragraph('页面需要在生成前明确显示预计结果数量；生成过程中显示任务进度；生成后支持单张查看、重试、选择、导出和继续编辑。')
doc.add_heading('六、生成后的继续编辑',1)
doc.add_paragraph('用户可以从结果中选择一张满意样片作为新的编辑基础。局部编辑只处理指定区域或指定对象，不重新理解全部人物、背景、道具和长提示词。')
for x in ['选择一张结果图作为当前版本。','选择编辑类型：背景、道具、服装、表情、手部、脚部、光线或其他局部。','框选、涂抹或指定需要修改的区域。','输入简短要求并生成多个局部版本。','保留原图与编辑版本，支持继续迭代或导出。']: b(x)
doc.add_heading('七、客户确认项',1); doc.add_paragraph('请确认以下功能理解是否符合实际使用需求：')
for x in ['首页同时支持“逐张批处理”和“多图融合生成”两种目标图模式。','人物参考图是所有生成任务的主体来源。','背景照片可用于自动替换背景，道具照片可用于替换指定道具。','道具替换需要用户明确指定画面中的替换对象。','生成结果可以继续进行局部编辑。','样片数量按照当前任务模式计算，并在生成前明确显示预计数量。','生成结果支持查看、重试、筛选、导出和版本延续。']: b(x,True)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=f.add_run('首页图像创作工作台 · 功能确认说明'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(130,140,137)
doc.save(out); print(out)
